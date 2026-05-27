from docx import Document
import re
from datetime import datetime

import os, sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import shutil

import django
django.setup()
from django.conf import settings

from main.models import invoice

# Used in functions below and in views.py for the download page to order the filenames 
def file_num(filename):
        filename = os.path.basename(filename)
        if any([ext in filename for ext in ['.PDF', '.sh']]):
             return -1 # Don't want pandas trying to read invalid file types when writing invoices
        if filename[1].isdigit():
            return int(filename[0:2])
        return int(filename[0])

# Used in views.py for the home page to check whether the invoices have already been written for the latest remittance advice
def latest_filenum(path):
    latest_year = max([year for year in os.listdir(path) if len(year)==4])
    dir = os.path.join(path, latest_year)
    return max([file_num(file) for file in os.listdir(dir)])

# Used in views.py, draft_emails.py and invoices.py
def latest_filename(path):
    latest_year = max([year for year in os.listdir(path) if len(year)==4])
    dir = os.path.join(path, latest_year)
    return os.path.join(dir, max(os.listdir(dir), key=file_num))

def gather_sefton_remittance_advice(from_year='2013'):
    remittance_advice_files = []
    years = [year for year in sorted(os.listdir(settings.MEDIA_REMITTANCE)) if year >= from_year]
    for year in years:
        dir = settings.MEDIA_REMITTANCE / year
        filenames = sorted(os.listdir(dir), key=file_num)
        remittance_advice_files += [os.path.join(dir, filename) for filename in filenames if '.csv' in filename]
    return remittance_advice_files

def latest_invoice_batch():
    folder = latest_filename(settings.MEDIA_INVOICES)
    year = int(os.path.basename(os.path.dirname(folder)))
    batch_number = file_num(os.path.basename(folder))

    return list(invoice.objects.filter(
        year=year,
        batch_number=batch_number,
        obsolete=False,
    ).select_related('Resident').order_by('invoice_number'))


def invoice_path(invoice_):
    path = str(invoice_.filename)
    if os.path.isfile(path):
        return path

    latest_invoice_folder = latest_filename(settings.MEDIA_INVOICES)
    fallback_path = os.path.join(
        latest_invoice_folder,
        f'{invoice_.invoice_number} - {invoice_.Resident.name}.docx',
    )
    if os.path.isfile(fallback_path):
        return fallback_path

    raise FileNotFoundError(f'Could not find invoice file for invoice {invoice_.invoice_number}')

def invoice_period(invoice_):
    DATE_PATTERN = re.compile(r'\b\d{2}/\d{2}/\d{4}\b')

    dates = []
    document = Document(invoice_path(invoice_))

    for paragraph in document.paragraphs:
        dates += DATE_PATTERN.findall(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    dates += DATE_PATTERN.findall(paragraph.text)

    if not dates:
        return '', ''

    dates = [datetime.strptime(date, '%d/%m/%Y').date() for date in dates]
    return min(dates).strftime('%d %B %Y'), max(dates).strftime('%d %B %Y')

def gather_all_invoices_for_resident(first_name, last_name):
    valid_folder = lambda folder: not any(exception in folder for exception in ['.xlsx', '.zip', '.sh', 'OBSOLETE'])

    destination_path = settings.MEDIA_INVOICES / f'{first_name} {last_name} invoices'
    os.makedirs(destination_path, exist_ok=True)
    for year in sorted([year for year in os.listdir(settings.MEDIA_INVOICES) if len(year)==4]):
        invoice_batches = [folder for folder in os.listdir(settings.MEDIA_INVOICES/ year) if valid_folder(folder)]
        for invoice_batch in sorted(invoice_batches, key=file_num):
            file_path = os.path.join(settings.MEDIA_INVOICES, year, invoice_batch)
            files_to_copy = [file for file in os.listdir(file_path) if first_name in file and last_name in file]
            for file in files_to_copy:
                shutil.copyfile(os.path.join(file_path, file), os.path.join(destination_path, file))

# if __name__=='__main__':
#     gather_all_invoices_for_resident(first_name='', last_name='')