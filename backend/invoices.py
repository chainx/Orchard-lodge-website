import os
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from datetime import datetime
from docx import Document
from random import randint
import pandas as pd
from shutil import make_archive
import pathlib

from backend.get_sefton_data import get_remittance_advice
from backend.file_utils import latest_filename, file_num

import django
django.setup()
from django.conf import settings
from main.models import resident, invoice, sefton_payment

def main():
    # Test for invoice writing
    # debt = pd.DataFrame({'Amount': 100, 'PaymentItemDates' : '01/01/2025 - 28/01/2025', 'AdjustmentLabel' : ''}, index=[0])
    # write_inovice(settings.MEDIA_INVOICES, datetime.now().strftime("%d %B %Y"), resident.objects.get(id=100), debt, 10000, '015670', batch_number=10)

    get_remittance_advice()

    latest_remittance = latest_filename(settings.MEDIA_REMITTANCE)
    invoices, sefton_payments, batch_number, folder = get_invoice_data_from_sefton_csv(latest_remittance)

    # rewrite_invoice('inv_no', 'title first last', folder, invoices)
    
    compile_summary_table(invoices, sefton_payments, folder)
    write_invoices_and_update_db(invoices, sefton_payments, batch_number, folder, update_db=True)

#==================================================================================================================================================================

def write_inovice(folder, date, Resident, sub_items, total, invoice_number, batch_number=None):
    document = Document(settings.MEDIA_INVOICES_TEMPLATES / 'INVOICE TEMPLATE.docx')

    tbl = document.tables[0]
    tbl.cell(0,0).paragraphs[0].text = tbl.cell(0,0).paragraphs[0].text.replace("NAME OF RECIPIENT", Resident.name)
    tbl.cell(0,1).paragraphs[0].text = tbl.cell(0,1).paragraphs[0].text.replace('INV_NO', invoice_number)
    tbl.cell(0,1).paragraphs[1].text = tbl.cell(0,1).paragraphs[1].text.replace("RES_REF", Resident.customer_ref_no)
    tbl.cell(0,1).paragraphs[2].text = tbl.cell(0,1).paragraphs[2].text.replace("DATE TODAY", date)
    
    desc_col, amount_col = document.tables[1].cell(0,0), document.tables[1].cell(0,1)
    desc_idx_offset, amount_idx_offset = 7, 8
    desc_idx_offset_min, amount_idx_offset_min = 3, 4
    max_rows = 20

    if len(sub_items) + desc_idx_offset > len(desc_col.paragraphs):
        print(f'Invoice for {Resident.name} has too many subitems, may have to be edited manually')
        offset = len(sub_items) - max_rows
        if offset >= 0:
            desc_idx_offset, amount_idx_offset = desc_idx_offset_min+offset, amount_idx_offset_min+offset
        else:
            print(f'{-offset} additional rows have been added')
            desc_idx_offset, amount_idx_offset = desc_idx_offset_min, amount_idx_offset_min
            for n in range(-offset):
                desc_col.add_paragraph()
                amount_col.add_paragraph()

    count=0
    for debt in sub_items.itertuples():
        paragraph = desc_col.paragraphs[desc_idx_offset+count]
        reason=' ('+debt[3]+')' if debt[3]!='' and debt[3]!='MA' else ''
        paragraph.text = f'From {debt[2]}{reason}'
        
        paragraph =  amount_col.paragraphs[amount_idx_offset+count]
        paragraph.text = u'\u00A3' + f'{debt[1]:.2f}'

        count+=1

    for paragraph in document.tables[1].cell(1,1).paragraphs:
        if paragraph.text.count('TOTAL')==1: paragraph.text=paragraph.text.replace('TOTAL', f'{total/100:.2f}')

    document.save(os.path.join(folder, f'{invoice_number} - {Resident.name}.docx'))

def rewrite_invoice(invoice_number, resident_name, folder, invoices):
    invoice = [invoice for invoice in invoices if invoice['Resident'].name == resident_name][0]
    invoice['invoice_number'] = invoice_number
    write_inovice(folder, **invoice)

#==================================================================================================================================================================

def get_invoice_data_from_sefton_csv(filename, save_new_residents=True): # Obtains data pertinent to writing invoices and updating the database
    date, year = datetime.now().strftime("%d %B %Y"), datetime.now().strftime("%Y")
    batch_number = file_num(os.path.basename(filename))
    os.makedirs(os.path.join(settings.MEDIA_INVOICES, year), exist_ok=True)
    folder = os.path.join(settings.MEDIA_INVOICES, year, f'{batch_number}. {date}')
    invoice_number = max(invoice.objects.filter(obsolete=False).values_list('invoice_number', flat=True))
    invoice_number = "%05d" % (int(invoice_number)+1)

    df, payment_period = read_and_format_sefton_csv(filename)

    invoices, sefton_payments = [], []
    for index, df_row in df[['FormattedName', 'Sefton ID']].drop_duplicates().iterrows():
        
        res_name, sefton_id = df_row['FormattedName'], df_row['Sefton ID']
        res = get_or_add_resident(res_name, sefton_id, save=save_new_residents)

        invoice_total = 0 # Subtracted from Sefton cost to find total paid by Sefton
        if inv := compile_personal_contributions_and_sefton_payments(res, df, batch_number, date, invoice_number, cost_or_income=1):
            invoice_number = "%05d" % (int(invoice_number)+1)
            invoice_total = inv['total']
            invoices.append(inv)
        if payment := compile_personal_contributions_and_sefton_payments(res, df, batch_number, payment_period, cost_or_income=0, invoice_total=invoice_total):
            sefton_payments.append(payment)
    
    invoices += compile_invoices_for_private_residents(payment_period, batch_number, date, invoice_number)
    return invoices, sefton_payments, batch_number, folder

def compile_personal_contributions_and_sefton_payments(res, df, batch_number, date, invoice_number=None, cost_or_income=1, invoice_total=0):
    # cost_or_income = 1 for personal contributions and 0 for Sefton payments
    filt = (df['Sefton ID']==res.sefton_id) & (df['IsIncome']==cost_or_income)
    if not df.loc[filt].empty:
        if abs(df.loc[filt]['Amount'].sum())<1:
            return # Don't write invoices for less that £1
        row = {
            'Resident': res,
            'date' : date,
            'sub_items' : df.loc[filt][['Amount', 'PaymentItemDates', 'AdjustmentLabel']],
            'total' : round(df.loc[filt]['Amount'].sum()*100) - invoice_total, # Amounts are stored in pennies in the database
            'batch_number': batch_number
        }
        if invoice_number:
            row['invoice_number'] = invoice_number

        return row

def compile_invoices_for_private_residents(payment_period, batch_number, date, invoice_number):
    private_invoices = []
    residents = resident.objects.filter(current=True, private=True).order_by('last') # Add data for private residents
    for res in residents:
        invoice_number = "%05d" % (int(invoice_number)+1)
        row = {
            'Resident' : res,
            'date' : date,
            'sub_items' : pd.DataFrame({'Amount': res.private_rate, 'PaymentItemDates' : payment_period, 'AdjustmentLabel' : ''}, index=[0]),
            'total' : res.private_rate*100, # Amounts are stored in pennies in the database
            'batch_number': batch_number,
            'invoice_number' : invoice_number
        }
        private_invoices.append(row)
    return private_invoices

#=============================================   UPDATE DB   ====================================================================================================

# Extract arguments for updating database from arguments for writing invoice, creates resident instance for new residents
def extract_invoice_args_for_db(invoice_data, folder):
    invoice_data['filename'] = os.path.join(folder, f"{invoice_data['invoice_number']} - {invoice_data['Resident']}.docx")
    invoice_data['date'] = datetime.strptime(invoice_data['date'], '%d %B %Y')
    invoice_data['year'] = invoice_data['date'].year
    invoice_data.pop('sub_items')
    return invoice_data

def extract_payment_args_for_db(payment):
    payment['date'] = datetime.strptime(payment['date'].split(' to ')[1], '%d/%m/%Y').date()
    payment['year'] = payment['date'].year
    payment.pop('sub_items')
    return payment

def write_invoices_and_update_db(invoices, sefton_payments, batch_number, folder, update_db=True):    
    year = datetime.now().strftime("%Y")
    if invoice.objects.filter(year=year, batch_number=batch_number):
        raise ValueError(f'There already exist entries in the database for batch number {batch_number} in {year}')

    pathlib.Path(folder).mkdir(parents=True, exist_ok=True)

    for invoice_data in invoices:
        write_inovice(folder, **invoice_data) # Write invoice file locally
        if update_db:
            invoice(**extract_invoice_args_for_db(invoice_data, folder)).save() # Save invoices to the database
    if update_db:
        for payment in sefton_payments:
            sefton_payment(**extract_payment_args_for_db(payment)).save()

    make_archive(folder, "zip", folder)

def compile_summary_table(invoices, sefton_payments, folder, print_table=False):
    invoices_df = pd.DataFrame(invoices, columns=['Resident', 'total'])
    payments_df = pd.DataFrame(sefton_payments, columns=['Resident', 'total'])
    df = pd.merge(invoices_df, payments_df, on='Resident', how='outer')
    df = df.rename(columns={'total_x': 'Invoice total', 'total_y': 'Sefton contribution'})
    df[['Invoice total', 'Sefton contribution']] /= 100
    df.index = df.reset_index(drop=True).index + 1
    df.to_excel(folder+'.xlsx')
    if print_table:
        print(df)

#=============================================   UTILS   =======================================================================================================

def read_and_format_sefton_csv(filename):
    try:
        df = pd.read_csv(filename)
    except:
        df = pd.read_csv(filename, encoding='ISO-8859-1')
    
    df = df.fillna('')

    unusual_rows = df[df['ServiceTotalLabel']!='Total for Orchard Lodge Care Home']
    if not unusual_rows.empty:
        df = df[df['ServiceTotalLabel']=='Total for Orchard Lodge Care Home'] # Remove unusual rows

    payment_period = df.iloc[0]['ReportContext'].split('Payment Period from ')[1]

    format_function = lambda row: row['Person'].split()[0] + ',' + ', '.join(row['ClientName'].split(' (')[0].split(',')[::-1])
    df['FormattedName'] = df.apply(format_function, axis=1)
    df['Sefton ID'] = df['ClientName'].apply(lambda x: x.split(' (')[1][:-1])
    
    return df, payment_period

def get_or_add_resident(res_name, sefton_id=None, current=True, private=False, save=True):
    new_res_created = False
    try:
        res = resident.objects.get(sefton_id=sefton_id)
    except:
        title, first, last = res_name.split(', ')
        if res := get_residents_with_similar_name(first, last):
            if not res.sefton_id:
                res.sefton_id = sefton_id
                res.save()
        else:
            new_res_created = True
            res = resident(title=title, first=first, last=last, sefton_id=sefton_id)
            res.current = current
            res.private = private
            res.customer_ref_no = generate_unique_customer_reference_number()
            if save:
                print(f'New resident added: {res.name} - {res.customer_ref_no}\n')
                res.save()
            else:
                print(f'New resident: {res.name} - {res.customer_ref_no}, NOT SAVED!\n')
    if res.private and not new_res_created:
        print(f'The private resident {res.name} has been added to the Sefton remittance advice\n')
    if not res.current and not new_res_created:
        print(f'The former resident {res.name} has been added to the Sefton remittance advice\n')
    return res

def get_residents_with_similar_name(first, last):
    # Create a regex pattern that matches with any number of spaces or apostrophes between characters
    regex_pattern = r'(?i)^' + r'[\s\']*'.join(last.replace(' ','').replace("'", '')) + r'[\s\']*$'
    close_matches = resident.objects.filter(first__istartswith=first.split()[0], last__regex=regex_pattern)
    if len(close_matches)==1:
        print(f'The resident {first} {last} has been identified with the following close match:\n{close_matches[0]}\n')
        return close_matches[0]
    elif len(close_matches)>1:
        raise ValueError(f'Multiple close matches for {first} {last} in the database:\n {close_matches}\n')

def generate_unique_customer_reference_number():
    while True:
        customer_ref_no = ''.join([str(randint(0, 9)) for n in range(6)])
        if not resident.objects.filter(customer_ref_no=customer_ref_no):
            break # Keep generating customer reference numbers until a unique one is generated
    return customer_ref_no

#==================================================================================================================================================================

if __name__=='__main__':
    main()
