import os, shutil
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from datetime import datetime, date
import pandas as pd
import sqlite3
from docx import Document
from pypdf import PdfWriter, PdfReader
from copy import deepcopy
import re
import subprocess
from shutil import which

import django
django.setup()
from django.conf import settings
from main.models import global_variables, resident, invoice, payment
from backend.invoices import get_residents_with_similar_name

BASE_FOLDER = settings.MEDIA_INVOICES / 'Statements of account'

def main():
    # write_cover_letter(BASE_FOLDER,1,resident.objects.get(id=88)) #193
    # write_statement_of_account(BASE_FOLDER,1,resident.objects.get(id=198))
    
    # produce_resident_table(first_name, last_name)

    produce_statements_of_account()
    # produce_statements_of_account(recently_left=True)

#==================================================================================================================================================================

def produce_statements_of_account(recently_left=False, include_cover_letter=True, convert_to_pdf=None):    
    auto_convert_to_pdf = convert_to_pdf is None
    recently_left_cutoff_date = global_variables.load().recently_left_cutoff_date
    if convert_to_pdf is None:
        convert_to_pdf = can_convert_to_pdf()

    filename = f"Statements of Account ({datetime.now().date().strftime('%d-%m-%Y')}){' Recently Left' if recently_left else ''}"
    folder = BASE_FOLDER / filename
    if not recently_left:
        residents = resident.objects.filter(current=True)
    else:
        residents = resident.objects.filter(leave_date__gt=recently_left_cutoff_date)
    resident_list = sorted(list(residents), key = lambda x: x.total_owed(), reverse=True)

    os.makedirs(folder, exist_ok=True)
    
    count = 1
    for res in resident_list:
        if res.total_owed() != 0:
            try:
                statement_of_account = write_statement_of_account(folder, count, res, convert_to_pdf=convert_to_pdf)
                cover_letter = None
                if include_cover_letter:
                    cover_letter = write_cover_letter(folder, count, res, convert_to_pdf=convert_to_pdf)
                if convert_to_pdf:
                    merge_pdfs(cover_letter, statement_of_account)
            except RuntimeError as exc:
                if not auto_convert_to_pdf or not convert_to_pdf:
                    raise
                print(f'PDF conversion failed; falling back to DOCX statements: {exc}')
                convert_to_pdf = False
                write_statement_of_account(folder, count, res, convert_to_pdf=convert_to_pdf)
                if include_cover_letter:
                    write_cover_letter(folder, count, res, convert_to_pdf=convert_to_pdf)
            count+=1
        else:
            print(res.name)

    os.makedirs(folder/'docx files', exist_ok=True)
    for file in os.listdir(folder):  
        if (folder/file).is_file() and (file.endswith(".docx") or "Cover Letter" in file):
            shutil.move(folder/file, folder/'docx files'/file)

    shutil.make_archive(folder, "zip", folder)

    # produce_summary_table(resident_list, filename, recently_left)

def write_cover_letter(folder, count, res, convert_to_pdf=True):
    thresholds = global_variables.load().statement_cover_letter_thresholds
    cover_letter_lower_threshold = thresholds['lower']
    cover_letter_upper_threshold = thresholds['upper']
    urgent_cover_letter_upper_threshold = thresholds['urgent_upper']

    if res.total_owed() < cover_letter_lower_threshold:
        return
    elif res.total_owed() >= cover_letter_lower_threshold and res.total_owed() <= cover_letter_upper_threshold:
        document = Document(settings.MEDIA_INVOICES_TEMPLATES / 'STATEMENT OF ACCOUNT COVER LETTER TEMPLATE (1).docx')
    elif res.total_owed() < urgent_cover_letter_upper_threshold:
        document = Document(settings.MEDIA_INVOICES_TEMPLATES / 'STATEMENT OF ACCOUNT COVER LETTER TEMPLATE (2).docx')
    else:
        document = Document(settings.MEDIA_INVOICES_TEMPLATES / 'STATEMENT OF ACCOUNT COVER LETTER TEMPLATE (3).docx')

    text_fields = {
        "RES_NAME_1": f'{res.title} {res.first[0]} {res.last}',
        "RES_NAME_2": f'{res.title} {res.last}',
        "[AMOUNT]": str_total(res.total_owed()),
        "[DATE]": datetime.now().date().strftime('%-d %B %Y'),
    }

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for text_field, text_input in text_fields.items():
                if text_field in run.text:
                    run.text = run.text.replace(text_field, text_input)

    filename = os.path.join(folder, f'{count}. {res.name} - Statement of Account Cover Letter')
    document.save(f'{filename}.docx')
    if convert_to_pdf:
        docx_to_pdf(filename)
        return f'{filename}.pdf'
    return f'{filename}.docx'

def write_statement_of_account(folder, count, res, convert_to_pdf=True):
    cutoff_date = global_variables.load().payments_invoices_cutoff_date

    document = Document(settings.MEDIA_INVOICES_TEMPLATES / 'STATEMENT OF ACCOUNT TEMPLATE.docx')

    text_fields = {
        "RES_NAME": res.name,
        "RES_REF": res.customer_ref_no,
        "DATE_TODAY": datetime.now().date().strftime('%d/%m/%Y'),
    }
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for text_field, text_input in text_fields.items():
                if text_field in run.text:
                    run.text = run.text.replace(text_field, text_input)

    document.tables[0].cell(1, 0).text = str_total(res.total_invoiced())
    document.tables[0].cell(1, 1).text = str_total(res.total_payed())
    run = document.tables[0].cell(1, 2).paragraphs[0].add_run(str_total(res.total_owed()))
    run.bold = True

    for i, inv in enumerate(res.invoice_set.filter(obsolete=False).order_by('date').filter(date__gte=cutoff_date)):
        new_row = document.tables[1].add_row()
        copy_row_format(document.tables[1].rows[0], new_row)
        new_row.cells[0].text = inv.date.strftime('%d/%m/%Y')
        new_row.cells[1].text = inv.invoice_number
        new_row.cells[2].text = str(inv.batch_number)
        new_row.cells[3].text = str_total(inv.total)

    for i, pay in enumerate(res.payment_set.order_by('date').filter(date__gte=cutoff_date)):
        new_row = document.tables[2].add_row()
        copy_row_format(document.tables[2].rows[0], new_row)
        new_row.cells[0].text = pay.date.strftime('%d/%m/%Y')
        new_row.cells[1].text = str_total(pay.amount)
        new_row.cells[2].text = 'Bank transfer' if pay.type in ['Santander', 'RBS'] else pay.type

    filename = os.path.join(folder, f'{res.name} - Statement of Account')
    document.save(f'{filename}.docx')
    if convert_to_pdf:
        docx_to_pdf(filename)
        return f'{filename}.pdf'
    return f'{filename}.docx'

def produce_summary_table(resident_list, filename, recently_left):
    owed_col_name = f"Total Owed ({datetime.now().date().strftime('%d %B')})"
    if not recently_left:
        df = pd.DataFrame([(res.name, res.total_owed()/100) for res in resident_list], columns=["Name", owed_col_name])
    else:
        df= pd.DataFrame([(res.name, res.total_owed()/100, res.leave_date) for res in resident_list], columns=["Name", owed_col_name, "Leave Date"])

    previous_filename = next(file for file in os.listdir(BASE_FOLDER) if is_prev_excel_sheet(file, recently_left))
    previous_df = pd.read_excel(BASE_FOLDER/previous_filename).iloc[:, 1:] # Removes index col
    
    combined_df = df.merge(previous_df, on="Name", how="left")
    combined_df.index = combined_df.index + 1

    with pd.ExcelWriter(BASE_FOLDER/f"{filename}.xlsx", engine="xlsxwriter") as writer:
        combined_df.to_excel(writer, sheet_name="Sheet1")
        worksheet = writer.sheets["Sheet1"]

        base_fmt   = writer.book.add_format({"font_size": 13})
        pound_fmt  = writer.book.add_format({"num_format": "£#,##0.00", "font_size": 13})
        green_fmt  = writer.book.add_format({"num_format": "£#,##0.00", "font_color": "green", "font_size": 13})


        for col_idx, col in enumerate(combined_df.columns):
            col_idx+=1
            max_len = max(
                combined_df[col].astype(str).map(len).max(),
                len(str(col))
            )
            if pd.api.types.is_numeric_dtype(combined_df[col]):
                worksheet.set_column(col_idx, col_idx, max_len, pound_fmt)

                worksheet.conditional_format(
                    1, col_idx, len(combined_df), col_idx,
                    {
                        "type": "cell",
                        "criteria": "<=",
                        "value": 0,
                        "format": green_fmt,
                    }
                )
            else:
                worksheet.set_column(col_idx, col_idx, max_len, base_fmt)

def produce_resident_table(first, last):
    res = get_residents_with_similar_name(first, last, verbose=False)

    conn = sqlite3.connect(settings.MEDIA_ROOT / "db.sqlite3")

    query = f"""
        SELECT date, amount, 'payment' AS type
        FROM main_payment
        WHERE Resident_id = {res.id}

        UNION ALL

        SELECT date, total, 'invoice' AS type
        FROM main_invoice
        WHERE Resident_id = {res.id}

        ORDER BY date
    """
    df = pd.read_sql_query(query, conn)
    conn.close()

    df["amount"] /= 100
    df["signed_amount"] = df["amount"].where(df["type"] == "invoice", -df["amount"])
    df["total_owed"] = df["signed_amount"].cumsum()
    df = df.drop(columns=["signed_amount"])
    df = df[["date", "type", "amount", "total_owed"]]

    filename = settings.MEDIA_INVOICES / f"{first} {last} summary table.xlsx"
    with pd.ExcelWriter(filename, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Sheet1")
        worksheet = writer.sheets["Sheet1"]
        worksheet.set_column(0, 1, 12, writer.book.add_format({"font_size": 13}))
        worksheet.set_column(2, 3, 12, writer.book.add_format({"num_format": "£#,##0.00", "font_size": 13}))

        chart = writer.book.add_chart({"type": "line"})

        chart.add_series({
            "name": "Total owed",
            "categories": ["Sheet1", 1, 0, len(df), 0],  # dates in col A
            "values":     ["Sheet1", 1, 3, len(df), 3],  # total_owed in col D
        })

        chart.set_title({"name": "Total Owed Over Time"})
        chart.set_x_axis({"name": "Date", "date_axis": True})
        chart.set_y_axis({"name": "Total Owed", "num_format": "£#,##0.00"})
        chart.set_legend({"none": True})

        worksheet.insert_chart("F2", chart)

def copy_row_format(source_row, target_row):
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):        
        tc_pr = source_cell._element.get_or_add_tcPr()
        new_tc_pr = target_cell._element.get_or_add_tcPr()
        for child in tc_pr.iterchildren():
            new_tc_pr.append(deepcopy(child))

def str_total(total):
    return u'\u00A3' + f'{total/100:,.2f}'

def is_prev_excel_sheet(file, recently_left):
    if '.xlsx' in file and 'Statements of Account' in file and '~lock' not in file:
        if not recently_left and 'Recently Left' not in file:
            return True
        if recently_left and 'Recently Left' in file:
            return True

def can_convert_to_pdf():
    return which('soffice') is not None

def docx_to_pdf(input_file):
    # Requires LibreOffice
    if not can_convert_to_pdf():
        raise RuntimeError('LibreOffice command "soffice" is not available; cannot convert DOCX to PDF.')

    cmd = [
        "soffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", os.path.dirname(input_file),
        f'{input_file}.docx'
    ]
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    pdf_file = f'{input_file}.pdf'
    if result.returncode != 0 or not os.path.exists(pdf_file):
        raise RuntimeError(
            f'LibreOffice failed to convert {input_file}.docx to PDF: '
            f'{result.stderr.decode(errors="replace")}'
        )

def merge_pdfs(cover_letter, statement_of_account):
    writer = PdfWriter()
    for pdf in [cover_letter, statement_of_account]:
        if pdf is None:
            continue
        reader = PdfReader(pdf)
        for page in reader.pages:
            writer.add_page(page)
    with open(statement_of_account, "wb") as f:
        writer.write(f)

#==================================================================================================================================================================

if __name__=='__main__':
    main()
