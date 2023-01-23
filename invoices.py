from datetime import datetime
from docx import Document
import pandas as pd
from shutil import make_archive
import os

from web_scraping import scrape_data

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'OrchardLodge.settings')
import django
django.setup()
from main.models import resident, invoice

def write_inovice(folder,date,Resident,sub_items,total,invoice_number):

    document = Document('Invoices/INVOICE TEMPLATE.DOCX')

    for paragraph in document.paragraphs:
        if paragraph.text.count("NAME OF RECIPIENT")==1: paragraph.text=paragraph.text.replace("NAME OF RECIPIENT",Resident)
        if paragraph.text.count("DATE TODAY")==1: paragraph.text=paragraph.text.replace("DATE TODAY",date)
        if paragraph.text.count("INVOICE NUMB")==1: paragraph.text=paragraph.text.replace("INVOICE NUMB",invoice_number)

    count=0
    for debt in sub_items.itertuples():
        paragraph = document.tables[0].cell(0,0).paragraphs[7+count]

        reason=''
        if debt[3]!='' and debt[3]!='MA': reason=' ('+debt[3]+')'
        paragraph.text='From ' + debt[2] + reason
        
        paragraph =  document.tables[0].cell(0,1).paragraphs[8+count]
        paragraph.text=chr(163)+str(debt[1])
        count+=1

    for paragraph in document.tables[0].cell(1,1).paragraphs:
        if paragraph.text.count('TOTAL')==1: paragraph.text=paragraph.text.replace('TOTAL',str(total))

    document.save(folder+invoice_number+' - '+Resident+'.docx')


#==============================================================================================================================================================================


def db(kwargs): #Extract arguments for updating database from arguments for writing invoice, creates resident instance for new residents
    name = kwargs['Resident'].split()
    kwargs['filename'] = kwargs['folder']+kwargs['invoice_number']+' - '+kwargs['Resident']+'.docx'
    kwargs['Resident'] = resident.objects.get_or_create(title=name[0], first=' '.join(name[1:-1]), last=name[-1])[0]
    kwargs['date'] = datetime.strptime(kwargs['date'],'%d %B %Y').strftime('%Y-%m-%d')
    kwargs['total'] *= 100 #Amounts are stored in pennies in the database
    [kwargs.pop(key) for key in ['folder','sub_items']]
    return kwargs

def handle_data(filename): #Obtains data pertinent to writing invoices and updating the database
    date, year = datetime.now().strftime("%d %B %Y"), datetime.now().strftime("%Y") 
    folder = 'Invoices/'+year+'/'+str(len(os.listdir('invoices/'+year))//2+1)+'. '+ date + '/'
    with open('Invoices/Invoice number.txt', 'r') as file:
        invoice_number = file.readline() 

    df = pd.read_csv(filename).fillna('') #if ServiceTotalLabel != 'Total for Orchard Lodge Care Home' then it may contain a note from Sefton
    payment_period = df.iloc[0]['ReportContext'].split('Payment Period from ')[1]

    kwargs_list=[]
    for res in df.Person.unique():
        filt = (df['Person']==res) & (df['IsIncome']==1)
        if not df.loc[filt].empty:
            kwargs = {
                'folder' : folder,
                'Resident' : res,
                'date' : date,
                'sub_items' : df.loc[filt][['Amount','PaymentItemDates','AdjustmentLabel']],
                'total' : df.loc[filt]['Amount'].sum(),
                'invoice_number' : invoice_number
            }
            kwargs_list.append(kwargs)
            invoice_number = "%05d" % (int(invoice_number)+1)

    residents = resident.objects.filter(current=True,private=True).order_by('last') #Add data for private residents
    for res in residents:
        kwargs = {
            'folder' : folder,
            'Resident' : res.name,
            'date' : date,
            'sub_items' : pd.DataFrame({'Amount': res.private_rate, 'PaymentItemDates' : payment_period, 'AdjustmentLabel' : ''}, index=[0]),
            'total' : res.private_rate,
            'invoice_number' : invoice_number
        }
        kwargs_list.append(kwargs)
        invoice_number = "%05d" % (int(invoice_number)+1)
    
    return kwargs_list,invoice_number

def write_invoices_and_update_db(kwargs_list,invoice_number):
    #Creating necessary folders if they don't already exist
    date, year = datetime.now().strftime("%d %B %Y"), datetime.now().strftime("%Y")
    if year not in os.listdir('Invoices'):
        os.makedirs('Invoices/'+year)
    folder = 'Invoices/'+year+'/'+str(len(os.listdir('invoices/'+year))//2+1)+'. '+ date + '/'
    os.makedirs(folder) 
    
    for kwargs in kwargs_list:
        write_inovice(**kwargs) #Write invoice file locally
        invoice(**db(kwargs)).save() #Save invoices to the database

    with open('Invoices/Invoice number.txt', 'w') as file: #Update invoice number
        file.write(invoice_number)

    make_archive(folder,"zip",folder)



if __name__=='__main__':

    from utils import get_latest
    latest_remittance = get_latest('Remittance advice')
    args = handle_data(latest_remittance)

    write_invoices_and_update_db(*args)